import json
from docx import Document

def generate_docx_from_swagger(swagger_file, output_file):
    with open(swagger_file, 'r') as file:
        swagger_data = json.load(file)

    document = Document()
    document.add_heading(swagger_data['info']['title'], level=1)
    document.add_paragraph(f"Version: {swagger_data['info']['version']}")
    document.add_paragraph(swagger_data['info']['description'] or '')

    for path, methods in swagger_data['paths'].items():
        for method, details in methods.items():
            document.add_heading(f"{method.upper()} {path}", level=2)
            document.add_paragraph(details.get('summary', ''))
            document.add_paragraph(details.get('description', ''))

            if 'parameters' in details:
                document.add_heading("Parameters:", level=3)
                for param in details['parameters']:
                    param_name = param.get('name', 'N/A')
                    param_desc = param.get('description', 'No description')
                    param_type = param.get('type', 'N/A')
                    document.add_paragraph(f"- {param_name} ({param_type}): {param_desc}")

            if 'responses' in details:
                document.add_heading("Responses:", level=3)
                for status, response in details['responses'].items():
                    document.add_paragraph(f"- {status}: {response.get('description', 'No description')}")

    document.save(output_file)
    print(f"Documentation saved to {output_file}")

# Uso do script
swagger_file = './swagger/swagger_api_doc.json'
output_file = './swagger/API_Documentation.docx'
generate_docx_from_swagger(swagger_file, output_file)
